"""Inspect all source legislation files to understand their structure."""
import os
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

import openpyxl
import docx
import csv

LEG = Path(r"C:\2026\369 Alliance\1. Legislation")
WORKING = Path(r"C:\2026\369 Alliance\1. Working folder")

def inspect_xlsx(path: Path, max_rows=8):
    print(f"\n{'='*80}\nXLSX: {path.name}\n{'='*80}")
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    for sn in wb.sheetnames:
        ws = wb[sn]
        print(f"\n  Sheet: {sn} ({ws.max_row} rows x {ws.max_column} cols)")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= max_rows: break
            cells = [str(c)[:60] if c is not None else "" for c in row[:8]]
            print(f"    R{i+1}: {' | '.join(cells)}")
    wb.close()

def inspect_docx(path: Path):
    print(f"\n{'='*80}\nDOCX: {path.name}\n{'='*80}")
    d = docx.Document(path)
    print(f"  Paragraphs: {len(d.paragraphs)}")
    for i, p in enumerate(d.paragraphs[:80]):
        if p.text.strip():
            print(f"    P{i}: {p.text[:200]}")
    print(f"  Tables: {len(d.tables)}")
    for ti, t in enumerate(d.tables[:3]):
        print(f"   Table {ti}: {len(t.rows)} rows x {len(t.columns)} cols")
        for ri, r in enumerate(t.rows[:5]):
            cells = [c.text.strip()[:40] for c in r.cells]
            print(f"     R{ri}: {' | '.join(cells)}")

def inspect_csv(path: Path, max_rows=5):
    print(f"\n{'='*80}\nCSV: {path.name}\n{'='*80}")
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i >= max_rows: break
            print(f"  R{i}: {' | '.join([str(c)[:60] for c in row[:8]])}")

# Inspect everything
inspect_xlsx(LEG / "NCC Building Reference Guide-Updated with new links 21 October 2021.xlsx")
inspect_docx(LEG / "AS in use.docx")
inspect_xlsx(WORKING / "New Master_Defect Library iAuditor_r.xlsx")
inspect_csv(LEG / "NCC" / "NCC_2022_HP_allClause.csv", 10)
inspect_csv(LEG / "NCC" / "NCC Volume 2 Standards.csv", 10)
