"""Build the Strata Story Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1A, 0x1A, 0x2E)
GOLD = RGBColor(0xA6, 0x8A, 0x64)
AMBER = RGBColor(0xC0, 0x70, 0x40)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def add_title(text, size=28, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color
    return p


def add_subtitle(text, size=14, color=GOLD, italic=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.color.rgb = color
    return p


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = AMBER
    return p


def add_para(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead + ' ')
        r.font.bold = True
    p.add_run(text)


def add_rule_box(rule, who, law):
    tbl = doc.add_table(rows=3, cols=1)
    tbl.autofit = True
    items = [
        ('Rule in one line: ', rule, NAVY),
        ('Who pays: ', who, AMBER),
        ('The law: ', law, GOLD),
    ]
    for i, (label, body, color) in enumerate(items):
        cell = tbl.rows[i].cells[0]
        shade_cell(cell, 'F4F1EA')
        cell.text = ''
        para = cell.paragraphs[0]
        r1 = para.add_run(label)
        r1.font.bold = True
        r1.font.color.rgb = color
        r1.font.size = Pt(10)
        r2 = para.add_run(body)
        r2.font.size = Pt(10)
    doc.add_paragraph()


def add_table(headers, rows, header_fill='1A1A2E'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        shade_cell(cell, header_fill)
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri % 2 == 0:
                shade_cell(cell, 'FAF7F0')
    doc.add_paragraph()


# ===== COVER =====
add_title('THE STRATA STORY', size=36, color=NAVY)
add_subtitle("A simple guide for owners — what's mine, what's \"ours\", and who pays when something breaks", size=14, color=GREY, italic=True)
add_para('Based on NSW Strata Legislation · Building Commission NSW · 369 Alliance', italic=True, color=GOLD, size=10)
doc.add_paragraph()

add_h2('How to read this guide')
add_para('This is a story, not a law book. Every chapter answers ONE question that owners ask us all the time.')
add_para('At the end of every chapter you will find:')
add_bullet('Rule in one line.', bold_lead='•')
add_bullet('Who pays.', bold_lead='•')
add_bullet('The law behind it.', bold_lead='•')
add_para('If you only have 5 minutes, read Chapter 2 (the imaginary line) and Chapter 5 (the decision tree). Those two answer 80% of the calls we get.', italic=True)

doc.add_page_break()

# ===== CHAPTER 1 =====
add_h1('CHAPTER 1 — What is Strata, really?')
add_para('Imagine 50 families decide to live in the same building. Every family wants their own front door, their own kitchen, their own bedroom. But everybody has to share the lift, the roof, the driveway, the walls between apartments, and the pipes inside the walls.')
add_para('Now imagine each family also has to share the cost of fixing those shared things — forever.')
add_para('That is Strata.', bold=True)
add_para('When you buy an apartment in NSW, you are not buying just an apartment. You are buying two things at once:')
add_bullet('the box of air inside your apartment.', bold_lead='Your Lot —')
add_bullet('a slice of everything that is shared.', bold_lead='A share of the Common Property —')
add_para("Together with all the other owners, you automatically become a member of a club called the Owners Corporation. You did not sign up. You cannot resign. The day you got the keys, you joined.")
add_rule_box(
    'Buying an apartment means buying a private box + a share of a shared building.',
    'You pay levies (your share of the bills) every quarter, for as long as you own.',
    'Strata Schemes Management Act 2015 (SSMA) and Strata Schemes Development Act 2015 (SSDA).'
)

# ===== CHAPTER 2 =====
add_h1('CHAPTER 2 — The Imaginary Line (the most important thing in this guide)')
add_para('There is a line that runs around the inside of every apartment. You cannot see it, but it is on the Strata Plan that was registered when the building was created.')
add_bullet('= Lot Property = YOUR responsibility.', bold_lead='Everything on YOUR side of the line')
add_bullet("= Common Property = the Owners Corporation's responsibility.", bold_lead='Everything on THE OTHER side of the line')
add_para('This is the rule for every strata scheme registered after 1 July 1974 (which is almost every modern apartment).')

add_h2('What is on YOUR side of the line (Lot Property)')
for it in [
    'Light fittings, light switches, power points, blinds and curtains.',
    'Internal walls not shown on the Strata Plan.',
    'Floor finishes (carpet, floorboards, tiles you laid yourself).',
    'Kitchen cupboards, benchtops, ovens, dishwashers, your appliances.',
    'Tiles and waterproofing on internal walls not shown on the Strata Plan.',
    'Bath, basin, toilet, taps and any plumbing above the upper surface of the floor.',
]:
    add_bullet(it)

add_h2('What is on THE OTHER side of the line (Common Property)')
for it in [
    'The structure — floor slabs, ceilings, columns, beams.',
    'External windows, external doors, balcony doors.',
    'The original floor tiles, including the original tiles in the shower.',
    'The original waterproofing membrane under those tiles.',
    'Tiles and waterproofing on boundary walls.',
    'The structural cubic space — shared plumbing, voids hiding shared plumbing/wires, and any wiring outside your apartment or shared between apartments.',
]:
    add_bullet(it)

add_rule_box(
    'If it was built originally and serves more than one owner, it is Common Property.',
    'Lot side = the owner. Common side = all owners (through levies).',
    'Strata Schemes Development Act 2015, Schedule 2, plus your registered Strata Plan.'
)

# ===== CHAPTER 3 =====
add_h1('CHAPTER 3 — The Cast (who is who in your building)')
add_para('Strata has more characters than a Netflix series. Knowing who does what saves you years of frustration.')
add_table(
    ['Character', 'What they actually do', 'Can they decide on remedial works?'],
    [
        ['The Owner (you)', 'Owns one Lot + share of Common Property. Pays levies. Votes at meetings.', 'Only on what is inside your own Lot.'],
        ['The Owners Corporation (OC)', 'Legal "company" made up of ALL owners. Owns the Common Property.', 'YES — the only body that can sign a remediation contract.'],
        ['The Strata Committee', 'A small group of owners (3–9) elected to handle day-to-day decisions.', 'Some, within authority limits set by the OC. Big jobs go back to the OC.'],
        ['The Strata Manager', 'Licensed agent hired to run admin, levies, meetings, insurance.', 'NO. Cannot sign remedial works contracts without OC approval.'],
        ['The Building Manager', 'Hired to look after the physical building (cleaning, gardens, small repairs).', 'NO. Cannot sign remedial works contracts.'],
        ['NCAT (the Tribunal)', 'Resolves disputes the OC cannot fix internally.', 'Decides disputes; orders works in some cases.'],
        ['NSW Fair Trading', 'Free mediation for strata disputes before NCAT.', 'No — they mediate.'],
        ['Building Commission NSW', 'The regulator. Inspects works. Issues fines. Cancels practitioner registrations.', "They police; they don't decide your project."],
    ]
)
add_rule_box(
    'Only the Owners Corporation can authorise major remedial works. Not the strata manager. Not the building manager. Not one loud owner.',
    'OC pays from the Capital Works Fund (or by special levy).',
    'Strata Schemes Management Act 2015, sections 5, 8, 9 and 106.'
)

# ===== CHAPTER 4 =====
add_h1('CHAPTER 4 — The Rule Book (which laws actually apply)')
add_para('Eight laws control your apartment life. You do not need to read them — you just need to know which one to point at when something happens.')
add_table(
    ['When this happens…', 'The law you point at'],
    [
        ['Meetings, levies, by-laws, voting, repairs of common property', 'Strata Schemes Management Act 2015 (SSMA) + Reg 2016'],
        ['The strata plan itself, subdivision, ending the scheme', 'Strata Schemes Development Act 2015 (SSDA) + Reg 2016'],
        ['Designs and builders for waterproofing, cladding, fire, structure on Class 2 buildings', 'Design and Building Practitioners Act 2020 (DBP) + Reg 2021'],
        ['Defects and serious problems in residential apartment buildings', 'Residential Apartment Buildings (Compliance and Enforcement Powers) Act 2020 (RAB)'],
        ['Builder warranties (2-year minor / 6-year major)', 'Home Building Act 1989 (HBA)'],
        ['Whether the work needs council approval (DA, CDC, exempt)', 'Environmental Planning and Assessment Act 1979 + Codes SEPP 2008'],
        ["The strata manager's licence and conduct", 'Property and Stock Agents Act 2002'],
        ['Buying or selling the apartment', 'Conveyancing Act 1919'],
    ]
)
add_rule_box(
    'Day-to-day = SSMA. Building quality = DBP + RAB. Defects warranty = HBA. Approval needed? = EPA Act + Codes SEPP.',
    'Depends on the matter — see other chapters.',
    'See each Act above.'
)

# ===== CHAPTER 5 =====
add_h1('CHAPTER 5 — The Decision Tree (whose problem is it?)')
add_para('Use this BEFORE you call anyone. It will save you weeks of arguing.')

add_h2('The 3 questions that almost always resolve it')
add_bullet('(inside the line or outside?)', bold_lead='Where is it?')
add_bullet('(original = usually Common Property)', bold_lead='Was it there from day 1?')
add_bullet('(yes = Common Property)', bold_lead='Does it serve more than one apartment?')

add_h2('Quick answers')
add_table(
    ['Situation', 'Answer'],
    [
        ['LOT side of the line (light fitting, your dishwasher, tiles you renovated, plumbing above floor surface)', 'YOU pay. YOU fix.'],
        ['ORIGINAL to the building (original tiles, original waterproofing, structure, external window, shared pipe)', 'OWNERS CORP pays and fixes.'],
        ['Newer item that is not original (e.g. you installed it yourself)', 'Whoever caused it / installed it pays and fixes.'],
    ]
)
add_rule_box(
    'Original + serves more than one Lot = Owners Corporation. Yours alone, installed by you, above the floor = you.',
    'See above.',
    'SSMA 2015 s.106 (duty of OC to maintain common property) + Schedule 2 of SSDA 2015.'
)

# ===== CHAPTER 6 =====
add_h1('CHAPTER 6 — The Remediation Story')
add_para('Big = the roof leaks across 6 apartments. The façade is cracking. The basement waterproofing failed. The cladding has to come off.')
add_para('These are remedial building works and they have their own rulebook now (since 2020).')

steps = [
    ('Step 1 — Spot it', 'Anyone can spot it: an owner, the building manager, an inspector. The issue is reported to the Strata Committee or the Strata Manager.'),
    ('Step 2 — Get advice', 'If nobody knows whether council approval is needed, the OC engages a Planning Consultant (Planning Institute of Australia register). They will tell you: Exempt Development, Complying Development Certificate (CDC), or Development Application (DA).'),
    ('Step 3 — The Owners Corporation VOTES', 'Nothing happens until the OC passes a motion at a general meeting. The strata manager cannot sign. The building manager cannot sign. One owner cannot sign.'),
    ('Step 4 — The DBP twist (the most misunderstood rule in NSW)', "For Class 2 buildings, even works that don't need council approval might still need a Registered Design Practitioner + a Registered Building Practitioner."),
    ('Step 5 — Lodge designs on the NSW Planning Portal', 'The Registered Building Practitioner uploads regulated designs, design compliance declarations, and contractor insurances. Nothing on site can start before lodgement.'),
    ('Step 6 — Build', 'Build exactly to the lodged designs. Any change becomes a Variation — a new design + declaration is lodged on the Portal during the works.'),
    ('Step 7 — Close out', 'Registered Building Practitioner lodges Building Compliance Declaration + Variation Statement + Contractor Document. This is the only document that cannot be delegated. Then the case is closed.'),
]
for h, body in steps:
    add_h2(h)
    add_para(body)

add_h2('Two famous DBP exceptions that catch everyone out')
add_table(
    ['Type of work', 'Is DBP usually triggered?'],
    [
        ['Waterproofing', "YES. Unless ALL three apply: (1) it's due to bathroom/kitchen/laundry/toilet alteration, AND (2) only one single dwelling, AND (3) it is exempt development."],
        ['External cladding', 'YES. Unless the work is NOT being done because of a council order or development control order. If there is an order, DBP applies.'],
    ]
)
add_para('Even if your work is "exempt development" under the Codes SEPP, waterproofing and cladding still need a Registered Design Practitioner.', bold=True, color=AMBER)
add_rule_box(
    'Owners vote → designer designs → portal upload → builder builds → builder declares → case closed.',
    'OC pays (capital works fund or special levy).',
    'DBP Act 2020, DBP Regulation 2021 (clause 13 has the exclusion list), RAB Act 2020, and Codes SEPP 2008.'
)

# ===== CHAPTER 7 =====
add_h1('CHAPTER 7 — The "It\'s an Emergency" Exception')
add_para("Sometimes you cannot wait for designs. The roof is open. People can't sleep in their apartment. There is real risk to safety.")
add_para('This is Emergency Remedial Building Work. The rules bend — but only a little.')
add_para('The OC must engage a Registered Building Practitioner to confirm in writing that there is a "reasonable excuse". To qualify, ALL of these must be true:', bold=True)
for it in [
    'Immediate action is needed.',
    'The issue is causing or likely to cause damage.',
    'It impacts: ability to live in the building, OR health and safety, OR risk of further damage.',
    'The impact is serious.',
    'The work done is only what is necessary to mitigate it — not the full final fix.',
]:
    add_bullet(it)
add_para('If all five apply, the builder can start without a regulated design. But within 7 days of finishing the emergency work, they must lodge on the Portal: a registration of the emergency work + a Building Compliance Declaration.')
add_rule_box(
    'Emergency = act now, declare within 7 days, then plan the proper fix.',
    'OC pays (capital works fund or special levy).',
    'DBP Act 2020 — emergency remedial building work provisions.'
)

# ===== CHAPTER 8 =====
add_h1('CHAPTER 8 — The Penalty Story')
add_para('Owners often say: "Can we just skip the paperwork and get on with it?" No. Here is what the Building Commission NSW can do:')
add_table(
    ['What went wrong', 'Body corporate', 'Individual'],
    [
        ['Registered Design Practitioner failed to provide design compliance declarations', 'Up to $165,000', 'Up to $55,000'],
        ['Registered Building Practitioner failed to lodge declared designs', 'Up to $33,000', 'Up to $11,000'],
        ['Significant offences', 'Prosecution', 'Prosecution'],
    ]
)
add_para('On top of fines, the Building Commission NSW can issue:')
add_bullet('Stop Work Orders under the RAB Act.')
add_bullet('Building Work Rectification Orders — forcing the OC to undo and redo the work.')
add_bullet('Prohibition Orders — stopping the issue of the Occupation Certificate or strata plan registration.')
add_para('These orders sit on the public register. Future buyers will see them.', italic=True)
add_rule_box(
    'Cutting corners on DBP costs more than doing it properly.',
    'OC and/or practitioners.',
    'DBP Act 2020 + RAB Act 2020.'
)

# ===== CHAPTER 9 =====
add_h1('CHAPTER 9 — "But what about my renovation?"')
add_para('Renovations inside your Lot are governed by the SSMA 2015 (sections 108–110). The law puts your work into one of three buckets:')
add_table(
    ['Type', 'What it means', 'What you need'],
    [
        ['Cosmetic Work (s.109)', 'Painting, hanging pictures, curtains, replacing handles, installing a TV or hooks.', 'Nothing. Just go ahead.'],
        ['Minor Renovations (s.110)', 'New kitchen, built-in wardrobe, internal walls, recessed lights, hardwood/laminate flooring.', 'OC approval by ordinary resolution (50%+).'],
        ['Major Renovations', 'Anything that changes external appearance, touches structure, waterproofing, fire safety, or common property.', 'By-law by special resolution (75%+) — and usually a DBP-compliant designer + builder.'],
    ]
)
add_para('Trap: Replacing your bathroom waterproofing membrane is NOT a cosmetic job. It touches the original waterproofing — which is Common Property. You need OC approval AND a Registered Design Practitioner.', bold=True, color=AMBER)
add_rule_box(
    "If it changes the look from outside, touches structure, fire, or waterproofing — it's Major. Get a by-law and a registered practitioner.",
    'You pay for your renovation work; if it touches Common Property, the by-law often makes you responsible for ongoing maintenance.',
    'SSMA 2015 sections 108–110, plus DBP Act 2020 for waterproofing.'
)

# ===== CHAPTER 10 =====
add_h1('CHAPTER 10 — The First-Year Story (when the building is brand new)')
add_table(
    ['Moment', 'What happens'],
    [
        ['Registration of the Strata Plan', 'The Owners Corporation is born. The Developer holds all votes (they own all unsold lots).'],
        ['Initial Period', 'Until 1/3 of lots are sold to outside buyers. The Developer cannot make big decisions about Common Property.'],
        ['First AGM', 'Within 2 months of the end of the Initial Period. Owners elect the first Strata Committee. Developer hands over keys, plans, the Strata Roll, manuals, maintenance schedule.'],
        ['Building Bond (Class 2, 4+ storeys, not covered by HBC insurance)', 'Developer lodges 2% of the contract price with NSW Fair Trading.'],
        ['Inspections', 'Independent Building Inspector does an interim inspection (15–18 months) and final inspection (21–24 months).'],
        ['6-Year Major / 2-Year Minor Defects', 'Statutory warranties under the Home Building Act 1989 s.18B run from completion.'],
        ['DBP Statutory Duty of Care', '10-year duty of care to every owner — current and future — from designer, builder, manufacturer, supplier (DBP Act s.37). Retrospective to 11 June 2010.'],
    ]
)
add_rule_box(
    'Year 1–2 = developer hand-over and bond. Year 0–10 = DBP duty of care. Year 0–6 = HBA major defects warranty.',
    'Developer/builder for warranty work; OC for ongoing maintenance after handover.',
    'SSMA 2015 Part 3, SSDA 2015, DBP Act 2020 s.37, HBA 1989 s.18B and Part 11.'
)

# ===== CHAPTER 11 =====
add_h1('CHAPTER 11 — How to raise an issue (without making it personal)')
add_para('The law gives you a clear ladder. Climb one rung at a time.')
add_table(
    ['Step', 'What to do'],
    [
        ['1', 'Speak to your Strata Committee Secretary or your Strata Manager (in writing).'],
        ['2', 'Strata Committee tries to resolve it within their delegated authority.'],
        ['3', 'If unresolved, raise it as a Motion at the next General Meeting of the OC.'],
        ['4', 'If unresolved, request FREE mediation with NSW Fair Trading (most disputes require mediation BEFORE NCAT).'],
        ['5', 'If still unresolved, apply to NCAT (Strata & Community Schemes Division). NCAT can order works, money, by-laws.'],
    ]
)
add_rule_box(
    'Always escalate one step at a time. Mediation is free. NCAT is the final step.',
    'Each step has its own filing fee at NCAT.',
    'SSMA 2015 Part 12, NCAT Civil and Administrative Tribunal Act 2013.'
)

# ===== CHAPTER 12 =====
add_h1('CHAPTER 12 — The 7 Mistakes Owners Make')
add_table(
    ['#', 'The mistake', 'The fix'],
    [
        ['1', '"It\'s my apartment, I can do what I want."', 'Read your Strata Plan + by-laws first. Most things touch Common Property somewhere.'],
        ['2', '"The Strata Manager said it\'s fine, so it\'s fine."', 'The Strata Manager has NO authority to approve remedial works. Only the OC does. Get it in writing in meeting minutes.'],
        ['3', '"It\'s just waterproofing, we don\'t need a designer."', 'Waterproofing on common property = DBP territory. Skip a Registered Design Practitioner = up to $165,000 fine.'],
        ['4', '"We don\'t need a planning consultant, it\'s just a balcony."', 'Adding to a balcony, removing planters, replacing balustrades almost always needs DA approval AND DBP compliance.'],
        ['5', '"Let\'s pay the builder cash and skip the Portal."', 'No Portal lodgement = no Building Compliance Declaration = no closure. Future buyers will not touch the building.'],
        ['6', '"The defects are old, no point chasing now."', 'DBP statutory duty of care = 10 years. HBA major defects = 6 years. Both still alive for most buildings completed since 2010.'],
        ['7', '"We voted to fix it last year — that\'s enough."', 'An OC resolution older than 12 months may need to be re-voted. Always check before signing.'],
    ]
)

# ===== CHAPTER 13 =====
add_h1('CHAPTER 13 — The 30-Second Glossary')
glossary = [
    ('OC', 'Owners Corporation. The "company" of all owners.'),
    ('SC', 'Strata Committee. The 3–9 owners elected to do day-to-day work.'),
    ('SM / BM', 'Strata Manager / Building Manager. Hired professionals.'),
    ('Common Property', 'Everything outside the Lot line. Owned by all owners together.'),
    ('Lot Property', 'Everything inside the Lot line. Owned by you alone.'),
    ('By-law', 'A rule that binds every owner, registered against the Strata Plan.'),
    ('Strata Plan (SP)', 'The legal map showing every Lot and the Common Property.'),
    ('BMS / SMS', 'Building / Strata Management Statement. Governs shared facilities in mixed-use buildings.'),
    ('BMC', 'Building Management Committee. Manages shared facilities between strata schemes and other lots.'),
    ('DA / CDC / Exempt', 'Three planning approval pathways under the EP&A Act + Codes SEPP.'),
    ('DBP / RAB', 'The two 2020 Acts that brought practitioner registration and rectification powers.'),
    ('CIRD', 'Construction-Issued Regulated Design.'),
    ('DCD', 'Design Compliance Declaration.'),
    ('BCD', 'Building Compliance Declaration.'),
    ('PC / OC', 'Practical Completion / Occupation Certificate.'),
    ('DLP', 'Defects Liability Period (usually 12 months, contractual).'),
    ('NCAT', 'NSW Civil and Administrative Tribunal.'),
    ('NCC / BCA', 'National Construction Code / Building Code of Australia.'),
]
add_table(['Term', 'Meaning'], glossary)

doc.add_page_break()

# ===== 1-PAGE CHEAT SHEET =====
add_title('THE 1-PAGE CHEAT SHEET', size=24, color=NAVY)
add_subtitle('Print this and stick it on the fridge.', size=11, color=GREY)
add_h2('Whose problem is it?')
add_table(
    ['The thing', 'Lot or Common?', 'Who pays?'],
    [
        ['Light bulb in your kitchen', 'LOT', 'YOU'],
        ['Dishwasher leak', 'LOT', 'YOU'],
        ['Original bathroom tiles cracked', 'COMMON', 'OC'],
        ['New bathroom tiles you laid leak', 'LOT', 'YOU'],
        ['Original waterproofing failed', 'COMMON', 'OC'],
        ['External window is broken', 'COMMON', 'OC'],
        ['Internal wall fell', 'LOT (not on plan) / COMMON (on plan)', 'Check the plan'],
        ['Roof is leaking', 'COMMON', 'OC'],
        ['Pipe inside wall, serves only your unit', 'LOT', 'YOU'],
        ['Pipe inside wall, serves multiple units', 'COMMON', 'OC'],
        ['Front door of your apartment', 'COMMON (on the plan)', 'OC'],
        ['Lock on your front door', 'LOT (you can change with notice)', 'YOU'],
        ['Carpet in your apartment', 'LOT', 'YOU'],
        ['Carpet in the corridor', 'COMMON', 'OC'],
    ]
)
add_h2('Before you start ANY work, ask 4 questions')
add_bullet('Does it touch Common Property? → Need OC approval.')
add_bullet('Does it change appearance, structure, fire, or waterproofing? → Need a by-law (special resolution).')
add_bullet('Does it involve waterproofing or cladding on a Class 2 building? → Need a Registered Design Practitioner + Registered Building Practitioner.')
add_bullet('Does it need a DA or CDC? → Engage a Planning Consultant.')
add_para("If yes to ANY of 1–3 → it's not a DIY job.", bold=True, color=AMBER)

doc.add_paragraph()
add_para('Prepared by 369 Alliance · Sources: SSMA 2015, SSDA 2015, DBP Act 2020 & Reg 2021, RAB Act 2020, HBA 1989, Codes SEPP 2008, EP&A Act 1979, Building Commission NSW.', italic=True, color=GREY, size=9)

out = "C:/2026/369 Alliance/Website/docs/strata/storytelling/Strata_Story_for_Owners.docx"
doc.save(out)
print("Saved:", out)
